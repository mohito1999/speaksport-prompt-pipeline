from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_LINE_SPACING

ROOT = Path('/Users/mohitmotwani/Documents/SpeakSport New Customer Agent Prompting')
OUT = ROOT / 'output' / 'docx'
OUT.mkdir(parents=True, exist_ok=True)
DOCX = OUT / 'SpeakSport_San_Juan_Oaks_Transfer_Reduction_Diagnostic.docx'
ASSET = ROOT / 'tmp' / 'report_assets'
LOGO = ASSET / 'speaksport_logo.png'

GREEN = '008F45'; DARK = '006B3F'; MINT = 'E5F4EC'; PALE = 'F3F7F5'
INK = '17221D'; MID = '63706A'; GRID = 'D9E0DC'; GOLD = 'A97818'; WHITE = 'FFFFFF'
FONT = 'Arial'

doc = Document()
sec = doc.sections[0]
sec.page_width = Inches(8.5); sec.page_height = Inches(11)
sec.top_margin = Inches(0.72); sec.bottom_margin = Inches(0.72)
sec.left_margin = Inches(0.82); sec.right_margin = Inches(0.82)
sec.header_distance = Inches(0.32); sec.footer_distance = Inches(0.35)

def set_font(run, size=None, bold=None, italic=None, color=INK, name=FONT):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn('w:ascii'), name)
    run._element.get_or_add_rPr().rFonts.set(qn('w:hAnsi'), name)
    if size is not None: run.font.size = Pt(size)
    if bold is not None: run.bold = bold
    if italic is not None: run.italic = italic
    if color: run.font.color.rgb = RGBColor.from_string(color)
    return run

def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd'); tcPr.append(shd)
    shd.set(qn('w:fill'), fill)

def set_cell_margins(cell, top=100, start=130, bottom=100, end=130):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr(); tcMar = tcPr.first_child_found_in('w:tcMar')
    if tcMar is None:
        tcMar = OxmlElement('w:tcMar'); tcPr.append(tcMar)
    for m, v in [('top',top),('start',start),('bottom',bottom),('end',end)]:
        node = tcMar.find(qn('w:'+m))
        if node is None: node = OxmlElement('w:'+m); tcMar.append(node)
        node.set(qn('w:w'), str(v)); node.set(qn('w:type'), 'dxa')

def set_table_geometry(table, widths, indent=120):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tblPr = table._tbl.tblPr
    layout = tblPr.find(qn('w:tblLayout'))
    if layout is None: layout = OxmlElement('w:tblLayout'); tblPr.append(layout)
    layout.set(qn('w:type'), 'fixed')
    total = sum(widths)
    tblW = tblPr.find(qn('w:tblW'))
    if tblW is None: tblW = OxmlElement('w:tblW'); tblPr.append(tblW)
    tblW.set(qn('w:w'), str(total)); tblW.set(qn('w:type'),'dxa')
    tblInd = tblPr.find(qn('w:tblInd'))
    if tblInd is None: tblInd = OxmlElement('w:tblInd'); tblPr.append(tblInd)
    tblInd.set(qn('w:w'), str(indent)); tblInd.set(qn('w:type'),'dxa')
    grid = table._tbl.tblGrid
    for child in list(grid): grid.remove(child)
    for w in widths:
        gc = OxmlElement('w:gridCol'); gc.set(qn('w:w'), str(w)); grid.append(gc)
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            cell.width = Inches(widths[i]/1440)
            tcW = cell._tc.get_or_add_tcPr().find(qn('w:tcW'))
            if tcW is None: tcW = OxmlElement('w:tcW'); cell._tc.get_or_add_tcPr().append(tcW)
            tcW.set(qn('w:w'), str(widths[i])); tcW.set(qn('w:type'),'dxa')
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

def set_repeat_header(row):
    trPr = row._tr.get_or_add_trPr(); el = OxmlElement('w:tblHeader'); el.set(qn('w:val'),'true'); trPr.append(el)

def set_cell_border(cell, color=GRID, size='6'):
    tcPr = cell._tc.get_or_add_tcPr(); borders = tcPr.first_child_found_in('w:tcBorders')
    if borders is None: borders = OxmlElement('w:tcBorders'); tcPr.append(borders)
    for edge in ('top','left','bottom','right','insideH','insideV'):
        tag = 'w:'+edge; e = borders.find(qn(tag))
        if e is None: e = OxmlElement(tag); borders.append(e)
        e.set(qn('w:val'),'single'); e.set(qn('w:sz'),size); e.set(qn('w:color'),color)

def add_field(paragraph, field):
    r = paragraph.add_run(); fld = OxmlElement('w:fldSimple'); fld.set(qn('w:instr'), field); r._r.addnext(fld)

# Preset-derived styles: standard_business_brief with named SpeakSport brand overrides.
styles = doc.styles
normal = styles['Normal']; normal.font.name = FONT; normal.font.size = Pt(10.5); normal.font.color.rgb = RGBColor.from_string(INK)
normal._element.rPr.rFonts.set(qn('w:ascii'), FONT); normal._element.rPr.rFonts.set(qn('w:hAnsi'), FONT)
normal.paragraph_format.space_after = Pt(5); normal.paragraph_format.line_spacing = 1.12

for nm, size, before, after, color in [
    ('Title',30,0,8,DARK), ('Subtitle',13.5,0,16,MID),
    ('Heading 1',18,14,7,DARK), ('Heading 2',14,11,5,GREEN), ('Heading 3',11.5,8,3,DARK)]:
    st = styles[nm]; st.font.name = FONT; st.font.size = Pt(size); st.font.bold = nm != 'Subtitle'; st.font.color.rgb = RGBColor.from_string(color)
    st._element.rPr.rFonts.set(qn('w:ascii'), FONT); st._element.rPr.rFonts.set(qn('w:hAnsi'), FONT)
    st.paragraph_format.space_before = Pt(before); st.paragraph_format.space_after = Pt(after); st.paragraph_format.keep_with_next = True
    pPr = st._element.get_or_add_pPr()
    old_border = pPr.find(qn('w:pBdr'))
    if old_border is not None: pPr.remove(old_border)

for nm in ['List Bullet','List Number']:
    st = styles[nm]; st.font.name = FONT; st.font.size = Pt(10.3)
    st._element.rPr.rFonts.set(qn('w:ascii'), FONT); st._element.rPr.rFonts.set(qn('w:hAnsi'), FONT)
    st.paragraph_format.left_indent = Inches(0.5); st.paragraph_format.first_line_indent = Inches(-0.25)
    st.paragraph_format.space_after = Pt(4); st.paragraph_format.line_spacing = 1.12

for name, fill, color in [('Lead Callout',MINT,DARK),('Note Callout',PALE,INK),('Risk Callout','FFF5DF',GOLD)]:
    if name not in styles:
        st = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
    else: st = styles[name]
    st.font.name = FONT; st.font.size = Pt(10.5); st.font.color.rgb = RGBColor.from_string(color)
    st._element.rPr.rFonts.set(qn('w:ascii'), FONT); st._element.rPr.rFonts.set(qn('w:hAnsi'), FONT)
    st.paragraph_format.left_indent = Inches(0.15); st.paragraph_format.right_indent = Inches(0.15)
    st.paragraph_format.space_before = Pt(5); st.paragraph_format.space_after = Pt(8); st.paragraph_format.line_spacing = 1.12
    pPr = st._element.get_or_add_pPr(); shd = OxmlElement('w:shd'); shd.set(qn('w:fill'),fill); pPr.append(shd)
    borders = OxmlElement('w:pBdr'); left = OxmlElement('w:left'); left.set(qn('w:val'),'single'); left.set(qn('w:sz'),'22'); left.set(qn('w:color'), GREEN if name!='Risk Callout' else GOLD); borders.append(left); pPr.append(borders)

def add_para(text='', style=None, bold_prefix=None, align=None, color=None, size=None, italic=False, keep=False):
    p = doc.add_paragraph(style=style)
    if bold_prefix and text.startswith(bold_prefix):
        set_font(p.add_run(bold_prefix), size=size, bold=True, color=color or INK)
        set_font(p.add_run(text[len(bold_prefix):]), size=size, italic=italic, color=color or INK)
    else:
        set_font(p.add_run(text), size=size, italic=italic, color=color or INK)
    if align is not None: p.alignment = align
    if keep: p.paragraph_format.keep_with_next = True
    return p

def add_bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet'); p.paragraph_format.left_indent = Inches(0.5+0.25*level)
    set_font(p.add_run(text), size=10.3); return p

def heading(text, level=1): return doc.add_heading(text, level=level)

def page_break(): doc.add_page_break()

def add_picture(path, width=6.45, caption=None):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after = Pt(4)
    inline = p.add_run().add_picture(str(path), width=Inches(width))
    alt_map = {
        'call_outcomes.png': 'Distribution of 523 calls: 162 not marked forwarded, 267 forwarded to the golf shop, and 94 forwarded elsewhere.',
        'weekly_rate.png': 'Golf-shop transfer rate was 51.5 percent from 21 to 27 July and 50.8 percent from 28 July to 3 August.',
        'proshop_reasons.png': 'Bar chart of primary reasons for 267 golf-shop transfers, led by 120 immediate or unspecified routing requests.',
        'initiation_donut.png': 'Transfer initiation split: 153 caller-led immediately, 103 initiated by the AI, and 11 later caller escalations.',
        'opportunity_levers.png': 'Primary reduction levers for golf-shop transfers: 131 caller adoption, 76 product or workflow, 24 conversation UX, 21 knowledge and data, and 15 appropriate routing.',
    }
    alt = alt_map.get(Path(path).name, Path(path).stem.replace('_',' '))
    inline._inline.docPr.set('descr', alt); inline._inline.docPr.set('title', alt[:120])
    if caption:
        c = add_para(caption, color=MID, size=8.5, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        c.paragraph_format.space_before = Pt(4); c.paragraph_format.space_after = Pt(7)

def add_table(headers, rows, widths, header_fill=GREEN, font_size=9.2):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths)
    hdr = table.rows[0]; set_repeat_header(hdr)
    for i, h in enumerate(headers):
        shade(hdr.cells[i], header_fill); set_cell_border(hdr.cells[i])
        p = hdr.cells[i].paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.LEFT; p.paragraph_format.space_after = Pt(0)
        set_font(p.add_run(str(h)), size=font_size, bold=True, color=WHITE if header_fill==GREEN else DARK)
    for ridx, row in enumerate(rows):
        cells = table.add_row().cells
        for i, val in enumerate(row):
            if ridx % 2 == 1: shade(cells[i], 'F7F9F8')
            set_cell_border(cells[i]); p = cells[i].paragraphs[0]; p.paragraph_format.space_after = Pt(0); p.paragraph_format.line_spacing = 1.05
            set_font(p.add_run(str(val)), size=font_size, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table

def metric_strip(metrics):
    table = doc.add_table(rows=1, cols=len(metrics)); widths=[round(9360/len(metrics))]*len(metrics); widths[-1]+=9360-sum(widths)
    set_table_geometry(table,widths)
    for i,(big,label) in enumerate(metrics):
        cell=table.cell(0,i); shade(cell,GREEN); set_cell_border(cell,color='FFFFFF',size='5')
        p=cell.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(1)
        set_font(p.add_run(big),size=18,bold=True,color=WHITE); p.add_run('\n'); set_font(p.add_run(label),size=8.3,bold=True,color=WHITE)
    doc.add_paragraph().paragraph_format.space_after=Pt(2)

# Header and footer
header = sec.header
hp = header.paragraphs[0]; hp.alignment = WD_ALIGN_PARAGRAPH.LEFT; hp.paragraph_format.space_after = Pt(0)
logo_inline = hp.add_run().add_picture(str(LOGO), width=Inches(1.72))
logo_inline._inline.docPr.set('descr', 'SpeakSport logo'); logo_inline._inline.docPr.set('title', 'SpeakSport logo')
footer = sec.footer
ft = footer.paragraphs[0]; ft.alignment = WD_ALIGN_PARAGRAPH.RIGHT; ft.paragraph_format.space_before = Pt(4)
set_font(ft.add_run('SpeakSport · San Juan Oaks transfer reduction diagnostic · 4 August 2026  |  '), size=8, color=MID)
add_field(ft, 'PAGE')

# PAGE 1 - customer pack cover
p = add_para('CUSTOMER TRANSFER REDUCTION DIAGNOSTIC', color=GREEN, size=10.5, keep=True); p.paragraph_format.space_before=Pt(18); p.paragraph_format.space_after=Pt(3)
doc.add_paragraph('San Juan Oaks', style='Title')
doc.add_paragraph('Why calls are reaching the pro shop - and how to reduce avoidable transfers', style='Subtitle')
add_table(['Review period','Calls reviewed','Prepared'], [['21 July-4 August 2026','523','4 August 2026']], [3400,2560,3400], header_fill='E5F4EC', font_size=9.5)
add_para('Bottom line: the client’s perception is supported by the data.', style='Lead Callout', bold_prefix='Bottom line: ')
add_para('The AI forwarded 361 calls (69.0%). Of these, 267 went specifically to the golf shop/pro shop - 51.1% of all calls. The golf-shop transfer rate was essentially flat across the two complete seven-day periods (51.5% then 50.8%), so staff would not yet feel a tangible reduction.', size=11)
metric_strip([('523','calls reviewed'),('361','all forwards'),('267','to golf shop'),('51.1%','of all calls')])
heading('The central finding',1)
add_para('There is no single cause. Nearly half of golf-shop transfers are driven by callers asking for a person before stating a need; the remainder is concentrated in booking-service gaps, live information gaps, and avoidable interaction friction. The reduction plan therefore needs coordinated work across product, prompt/experience design, and course-supplied operational knowledge.', size=10.8)
add_para('This report is designed for joint review by SpeakSport and San Juan Oaks. It distinguishes transfers that are addressable from transfers that should remain human.', style='Note Callout')

# PAGE 2 Executive summary
page_break(); heading('Executive summary',1)
add_picture(ASSET/'call_outcomes.png', width=6.55, caption='Source: SpeakSport call export; a “forward” is identified from ended_reason = assistant-forwarded-call.')
heading('What the numbers say',2)
add_bullet('361 of 523 calls were forwarded (69.0%); 162 calls (31.0%) were not marked as forwarded. “Not forwarded” is not automatically equivalent to successfully resolved.')
add_bullet('267 calls reached the golf shop/pro shop - 74.0% of all forwards and 51.1% of every call in the sample.')
add_bullet('94 forwards (18.0% of all calls) went elsewhere, primarily to restaurants and events. These do not contribute directly to golf-shop load, but they explain why the overall forward rate is higher than the pro-shop rate.')
add_bullet('153 of the 267 golf-shop transfers (57.3%) began with an immediate caller request for a person or destination. 103 (38.6%) were initiated by the AI because of a gap/routing rule; 11 (4.1%) followed an attempted self-service interaction.')
add_bullet('The largest product-shaped cluster is booking servicing: 25 cancellations, 32 modifications/confirmations, and 9 group/multiple-tee-time requests reached the golf shop.')
add_para('Priority signal: the prior 31 July update states that on-call cancellations and upcoming-booking confirmations had been rolled out. Yet the reviewed calls still include 25 cancellation transfers and 32 modification/confirmation transfers, including a cancellation on 3 August. This should be treated first as a production/configuration/prompt consistency issue, not as a request for a brand-new capability.', style='Risk Callout', bold_prefix='Priority signal: ')

# PAGE 3 transfer landscape
page_break(); heading('1. Transfer landscape',1)
add_picture(ASSET/'weekly_rate.png', width=5.85, caption='Two complete seven-day periods; 4 August is excluded because it is a partial day in the export.')
add_para('Golf-shop transfer volume did not materially improve week over week. The overall forward rate rose from 66.5% to 71.8%, while the rate specifically reaching the golf shop stayed nearly unchanged: 51.5% versus 50.8%. Daily rates were volatile, so the two-week sample should be treated as a baseline rather than a long-term trend.', size=10.5)
heading('Where forwarded calls went',2)
dest_rows=[
    ('Golf shop / pro shop','267','74.0%','51.1%'),('36 Degrees North','36','10.0%','6.9%'),
    ("McCann’s Bar & Grill",'23','6.4%','4.4%'),('Events / group sales','17','4.7%','3.3%'),
    ('Other or named destinations','18','5.0%','3.4%')]
add_table(['Destination','Calls','% of forwards','% of all calls'],dest_rows,[4400,1400,1800,1760],font_size=9.3)
add_para('Interpretation: the customer’s operational concern is narrower than the overall 69.0% headline. The most relevant baseline is 51.1% of all calls reaching the golf shop.', style='Lead Callout', bold_prefix='Interpretation: ')

# PAGE 4 root cause distribution
page_break(); heading('2. Why calls reach the golf shop',1)
add_picture(ASSET/'proshop_reasons.png', width=6.55)
add_para('The 267 golf-shop transfers were assigned one primary intent. “Immediate / unspecified routing” means the caller asked for the golf shop, an operator, or a representative without revealing a resolvable task. “Other specialist needs” combines dining, administration, lessons, and events that happened to route to the golf shop.', size=9.5, color=MID)
heading('Who initiated the transfer?',2)
add_picture(ASSET/'initiation_donut.png', width=5.65)

# PAGE 5 immediate routing
page_break(); heading('3. Root cause A - callers ask for a person immediately',1)
add_para('153 golf-shop transfers (57.3%) were caller-led from the outset. Within the full golf-shop set, 120 calls (44.9%) contained no substantive need beyond “pro shop,” “golf shop,” “operator,” or an equivalent request.', style='Lead Callout')
heading('What is happening',2)
add_bullet('The caller arrives with a routing habit: they dial the main number expecting the pro shop, not an automated service.')
add_bullet('The assistant typically confirms the destination and transfers immediately. It does not consistently make a short, task-specific attempt to surface whether the need is one it can resolve.')
add_bullet('Some callers reveal a tee-time need only after asking for a person. These are the best candidates for a brief AI-first rescue; callers seeking a named staff member should still be routed promptly.')
heading('Recommended treatment',2)
add_para('Use a one-question soft gate, never a hard barrier:', bold_prefix='Use a one-question soft gate, never a hard barrier:')
add_para('“I can usually book, change, cancel, confirm, or price a tee time right here. Tell me what you need, or say ‘transfer’ and I’ll connect you.”', style='Lead Callout')
add_bullet('If the caller repeats the request, transfer immediately. Do not interrogate or require multiple confirmations.')
add_bullet('If the caller names a person, department, restaurant, emergency, or on-course issue, route without attempting containment.')
add_bullet('Measure the “rescue acceptance rate”: the share of immediate golf-shop requests that become a completed AI task after one prompt.')
add_para('Opportunity sizing: this lever is large but uncertain. Even a 10%-20% acceptance rate across the 131 golf-shop calls assigned to caller-adoption framing would avoid roughly 13-26 transfers in a comparable two-week period.', style='Note Callout', bold_prefix='Opportunity sizing: ')

# PAGE 6 capability mismatch
page_break(); heading('4. Root cause B - booking-service capability is not consistently reaching callers',1)
add_para('Booking servicing accounts for 66 golf-shop transfers: 25 cancellations, 32 modification/confirmation requests, and 9 group/multiple-tee-time requests. Lost-and-found adds another 11 calls that could be handled through a structured intake workflow rather than synchronous transfer.', size=10.8)
heading('The cancellation and confirmation inconsistency',2)
add_para('The 31 July performance update says on-call cancellations and upcoming-booking confirmations are live. The transcripts nevertheless contain repeated statements such as “I am not able to process tee-time cancellations” followed by transfer. This pattern appears throughout the review window and as late as 3 August.', style='Risk Callout')
add_bullet('Audit tenant feature flags, prompt versions, integration permissions, eligibility rules, and fallback branches for San Juan Oaks.')
add_bullet('Run scripted regression calls for cancel, confirm, change player count, change date/time, no-email confirmation, and duplicate booking scenarios.')
add_bullet('Log the exact reason every servicing attempt falls back: no matching reservation, identity mismatch, ineligible booking, API error, unsupported mutation, or prompt-only refusal.')
heading('Capability roadmap',2)
cap_rows=[
    ('P0','Cancel and confirm','Restore/verify the capability already described as live; remove stale “cannot cancel” branches.','25 cancels + part of 32 confirm/modify'),
    ('P0','Modify a booking','Support player-count, date, and time changes with eligibility checks and clear confirmation.','32 modify/confirm'),
    ('P1','Group requests','Capture group size/date/time and either book adjacent inventory or send a structured request to the correct owner.','9 group/multiple'),
    ('P1','Lost and found intake','Capture item, location, date/time, contact details, and asynchronously notify staff; transfer only for urgent same-day retrieval.','11 lost/found')]
add_table(['Priority','Capability','Action','Observed pool'],cap_rows,[900,1700,4700,2060],font_size=8.4)

# PAGE 7 knowledge and UX
page_break(); heading('5. Root cause C - live knowledge and interaction friction',1)
add_picture(ASSET/'opportunity_levers.png', width=6.55)
heading('Knowledge and live-data gaps',2)
add_para('21 golf-shop transfers were assigned primarily to knowledge/data integration. Common themes were exact green-fee pricing, replay or association rates, range closing times, maintenance/punching dates, twilight definitions, dress code, and current course conditions.', size=10.3)
add_bullet('Return the exact price alongside every available tee time. Avoid finding inventory and then transferring for the rate.')
add_bullet('Add a time-stamped operational feed for range closures, maintenance, frost/cart-path status, and same-day exceptions.')
add_bullet('For policies and promotions, store an answer plus an effective date, owner, and expiry/review date.')
heading('Conversation and booking UX friction',2)
add_para('24 golf-shop transfers were assigned primarily to conversation/UX improvement. The most visible friction appears in name/email capture, repeated clarification loops, silence/filler behavior, and booking flows that reach an available tee time but do not complete.', size=10.3)
add_bullet('Stop requiring NATO-phonetic spelling for routine names and email addresses. Confirm naturally, use the caller’s known profile/ANI when permitted, or send an SMS form/link for correction.')
add_bullet('Remove stray system utterances such as “I lost track,” “I’ll proceed with scheduling,” or repeated hold language that make the assistant feel unreliable.')
add_bullet('When confidence is low, summarize what was understood once and offer two choices. Do not keep asking open-ended clarification questions.')
add_bullet('Instrument booking drop-off by step: date, players, time, name, email, lookup, price, confirmation, payment/eligibility, and API completion.')

# PAGE 8 qualitative evidence
page_break(); heading('6. What callers experienced - representative patterns',1)
examples=[
    ('Immediate routing habit','Caller opens with “pro shop” or “operator”; the AI confirms and transfers without learning the task.','Introduce the one-question soft gate; honor a repeated transfer request.'),
    ('Cancellation refusal','Caller asks to cancel; the AI says it cannot process cancellations and routes to the golf shop.','Verify the already-announced cancellation capability and eliminate stale refusal paths.'),
    ('Confirmation / modification gap','Caller asks to confirm, add/remove a player, or change a time; the AI routes rather than servicing the reservation.','Add or restore authenticated lookup and mutation flows with explicit fallback reasons.'),
    ('Exact-price handoff','AI finds an available time but cannot state the exact fee, so it offers the golf shop.','Expose the booked-time rate in the same availability response.'),
    ('Identity capture frustration','A booking advances to spelling a name/email; repeated correction or phonetic requests lead the caller to ask for a person.','Use profile/ANI, natural confirmation, and SMS correction.'),
    ('Current-operations uncertainty','The AI knows normal range hours but transfers for a Monday or same-day exception.','Publish a time-stamped daily operations feed.'),
    ('Lost and found','The assistant has no record/workflow and transfers every inquiry.','Create structured intake and asynchronous staff notification.'),
]
add_table(['Pattern','Observed experience','Reduction approach'],examples,[1750,3740,3870],font_size=8.6)
add_para('Examples are paraphrased from reviewed transcripts and exclude caller names, phone numbers, and email addresses.', size=8.5, color=MID, italic=True)

# PAGE 9 joint action plan
page_break(); heading('7. Recommended joint action plan',1)
heading('First 14 days - stabilize and prove the basics',2)
actions1=[
    ('SpeakSport','Audit cancellation and confirmation production paths; publish failure-reason logging.','Owner assigned; regression suite passes; no stale “cannot cancel” prompt when eligible.'),
    ('SpeakSport','Deploy the one-question soft gate for immediate golf-shop requests with instant opt-out.','Rescue acceptance and repeat-transfer rates visible by day.'),
    ('SpeakSport','Remove phonetic-email requirement and stray filler/system utterances.','Lower transfer/drop-off after identity capture.'),
    ('San Juan Oaks','Provide the high-priority knowledge pack and named routing owners.','All P0 knowledge items have owner, effective date, and review cadence.'),
    ('Joint','Review 30 transferred calls after changes, using the taxonomy in this report.','Shared avoidable/appropriate labels; disputed calls resolved.')]
add_table(['Owner','Action','Definition of done'],actions1,[1500,4300,3560],font_size=8.7)
heading('Days 15-45 - expand self-service',2)
add_bullet('Enable exact tee-time pricing, booking modification, and richer reservation confirmation.')
add_bullet('Launch lost-and-found structured intake and group-booking intake with asynchronous routing.')
add_bullet('Add a daily operations feed for range hours, maintenance, course status, and temporary exceptions.')
add_bullet('Create transfer analytics by destination, initiating party, caller intent, fallback reason, and successful human connection.')
heading('Days 46-90 - optimize and govern',2)
add_bullet('Run controlled prompt experiments on the soft gate and booking recovery language.')
add_bullet('Set a monthly knowledge review with named San Juan Oaks owners and expiry alerts.')
add_bullet('Expand self-service only where the human workflow is stable, safe, and auditable; preserve direct transfer for emergencies, named staff, and specialist judgment.')

# PAGE 10 knowledge pack
page_break(); heading('8. San Juan Oaks knowledge and workflow pack',1)
add_para('The course can materially improve containment by supplying a small, governed set of information. Each item should include an owner, “last updated” date, and expiry/review cadence.', size=10.4)
kb=[
    ('Daily operations','Range opening/closing by day; Monday maintenance cutoff; frost/cart-path status; course closures; aeration/punching schedule; twilight start definition.','Daily or event-driven'),
    ('Pricing','Exact rate source; replay specials; NCGA/youth/ride-along policies; spectator/cart fees; range access; club rental; promotions and expiry.','Live feed or weekly'),
    ('Booking policy','Booking window; single-player rules; group threshold; multiple tee times; cancellation/change windows; no-show policy; late-arrival process.','Monthly / on change'),
    ('Facilities','Dress code; practice facilities; showers/locker rooms; accessibility; GPS/cart assistance; course amenities and hours.','Quarterly / on change'),
    ('Retail/admin','Gift card purchase and balance; merchandise; club fitting; receipts/refunds; membership contacts; donation/sponsorship/vendor requests.','Monthly'),
    ('Dining/events','Restaurant hours, menus and ordering routes; reservation platforms; patio/corkage policies; live music; banquets, weddings and tournaments.','Weekly / event-driven'),
    ('Lost and found','Who owns it; intake fields; response SLA; urgent vs non-urgent; where recovered items are stored; caller update method.','Quarterly / on change'),
    ('Routing map','Named staff/roles, extensions, hours, overflow destination, voicemail behavior, and which intents must bypass AI.','Monthly / staffing change'),
]
add_table(['Knowledge domain','Required content','Cadence'],kb,[1750,5900,1710],font_size=8.35)
add_para('Recommended operating rule: when information is volatile, the AI should state when it was last updated. If data is stale or unavailable, it should offer a structured callback/request before defaulting to a synchronous transfer.', style='Lead Callout', bold_prefix='Recommended operating rule: ')

# PAGE 11 scenarios
page_break(); heading('9. Reduction scenarios and measurement plan',1)
add_para('The scenarios below are planning models, not forecasts. They apply different “capture rates” to the observed golf-shop opportunity levers while leaving the 15 transfers judged appropriate for specialist routing unchanged.', size=10.2)
scenario=[
    ('Current baseline','-','0','267','51.1%','-'),
    ('Conservative','Adoption 10%; product 35%; UX 35%; knowledge 50%','59','208','39.8%','22%'),
    ('Target','Adoption 20%; product 60%; UX 60%; knowledge 75%','102','165','31.5%','38%'),
    ('Stretch','Adoption 30%; product 80%; UX 75%; knowledge 90%','137','130','24.9%','51%')]
add_table(['Scenario','Assumed capture by lever','Fewer golf-shop transfers','Remaining calls','Rate / all calls','Relative reduction'],scenario,[1300,3200,1500,1050,1050,1260],font_size=8.25)
add_para('A sensible near-term operating target is below 40 golf-shop transfers per 100 calls, followed by a 60-90 day target around 32 per 100 calls if capability and data work land successfully.', style='Lead Callout')
heading('Weekly scorecard',2)
score=[
    ('Golf-shop transfer rate','Golf-shop forwards / all calls','51.1%','<40% first; ~32% target'),
    ('Overall forward rate','All forwards / all calls','69.0%','Context only; destination matters'),
    ('Immediate human-request rate','Immediate caller-led golf-shop forwards / all calls','29.3%','Track; improve rescue acceptance'),
    ('AI-initiated golf-shop rate','AI-led golf-shop forwards / all calls','19.7%','Primary product/knowledge signal'),
    ('Servicing fallback rate','Cancel/modify/confirm forwards / relevant servicing requests','Not derivable from export','Instrument denominator and reason'),
    ('Rescue acceptance','Immediate human requests resolved after one soft gate / eligible requests','Not instrumented','10%-20% initial target'),
    ('Transfer connection success','Transfers answered / attempted','Not in export','Instrument by destination')]
add_table(['Metric','Definition','Baseline','Target / use'],score,[2100,3700,1650,1910],font_size=8.3)

# PAGE 12 methodology appendix
page_break(); heading('Appendix A - methodology and caveats',1)
heading('Data and classification',2)
add_bullet('Source file: calls-2026-07-21-to-2026-08-04.csv, containing 523 calls and transcript text.')
add_bullet('A transfer is defined strictly as ended_reason = assistant-forwarded-call. Destination is inferred from the final transfer statement in the transcript.')
add_bullet('Each forwarded call was reviewed using caller utterances and transfer context, then assigned one primary initiation type, intent, destination, and reduction lever. Categories are mutually exclusive for counting, even when a call contained more than one topic.')
add_bullet('The 4 August data is partial. Week-over-week comparison uses 21-27 July and 28 July-3 August only.')
add_bullet('Speech-to-text errors are visible (for example “tea time,” “key time,” or “Photoshop” for tee time/pro shop). Classification uses context, but isolated ambiguous calls remain possible.')
add_bullet('A call not marked forwarded may still be abandoned or unresolved. Resolution/containment should be measured separately once explicit outcome data is available.')
heading('All forwarded calls by primary intent',2)
intent_rows=[
    ('Immediate/unspecified routing','125','34.6%','23.9%'),('Dining/food service','68','18.8%','13.0%'),
    ('New tee time/availability','33','9.1%','6.3%'),('Booking modification/confirmation','32','8.9%','6.1%'),
    ('Cancellation','25','6.9%','4.8%'),('Membership/billing/retail/admin','17','4.7%','3.3%'),
    ('Group/multiple tee times','13','3.6%','2.5%'),('Lost and found','13','3.6%','2.5%'),
    ('Course/current conditions','12','3.3%','2.3%'),('Rates/promotions','10','2.8%','1.9%'),
    ('Events/tournaments','10','2.8%','1.9%'),('Lessons/other','3','0.8%','0.6%')]
add_table(['Primary intent','Calls','% of forwards','% of all calls'],intent_rows,[4300,1300,1850,1910],font_size=8.5)
heading('Sources',2)
add_para('1. SpeakSport call export for San Juan Oaks, 21 July-4 August 2026 (523 calls).\n2. SpeakSport Customer Performance Update - San Juan Oaks, prepared 31 July 2026.\n3. SpeakSport logo supplied 4 August 2026.', size=8.8, color=MID)

# Core properties
props=doc.core_properties; props.title='San Juan Oaks Transfer Reduction Diagnostic'; props.subject='Quantitative and qualitative analysis of call transfers'; props.author='SpeakSport'; props.keywords='SpeakSport, San Juan Oaks, transfers, AI receptionist, golf shop'

doc.save(DOCX)
print(DOCX)
