#!/usr/bin/env python3
"""Generate a branded, two-page SpeakSport customer performance update."""

from __future__ import annotations

import argparse
import json
import math
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


BRAND_GREEN = "008A45"
BRAND_DARK = "00633A"
INK = "17221C"
MUTED = "5F6B64"
PALE_GREEN = "EAF6EF"
PALE_GRAY = "F3F5F4"
WHITE = "FFFFFF"
CONTENT_WIDTH_DXA = 9360
CELL_MARGIN_DXA = {"top": 120, "bottom": 120, "start": 150, "end": 150}


def rgb(hex_value: str) -> RGBColor:
    return RGBColor.from_string(hex_value)


def set_run(run, *, size=None, bold=None, color=INK, italic=None, font="Calibri"):
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), font)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    run.font.color.rgb = rgb(color)


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, margins=CELL_MARGIN_DXA):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in margins.items():
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color="DDE5E0", size=6):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), str(size))
        node.set(qn("w:color"), color)


def set_table_geometry(table, widths_dxa: list[int], indent_dxa=120):
    total = sum(widths_dxa)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def remove_table_cell_paragraph_spacing(table):
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.0


def add_heading(doc, text: str, level=1):
    paragraph = doc.add_paragraph(style=f"Heading {level}")
    paragraph.add_run(text)
    return paragraph


def add_logo(paragraph, logo_path: Path, width: float):
    inline = paragraph.add_run().add_picture(str(logo_path), width=Inches(width))
    inline._inline.docPr.set("descr", "SpeakSport logo")
    inline._inline.docPr.set("title", "SpeakSport")


def add_metric_grid(doc, metrics):
    columns = 3 if len(metrics) > 4 else 2
    rows = math.ceil(len(metrics) / columns)
    widths = [CONTENT_WIDTH_DXA // columns] * columns
    table = doc.add_table(rows=rows, cols=columns)
    set_table_geometry(table, widths)
    set_table_borders(table, color="D7E4DC", size=5)
    for index, metric in enumerate(metrics):
        cell = table.rows[index // columns].cells[index % columns]
        set_cell_shading(cell, PALE_GRAY if index % 2 else PALE_GREEN)
        p_value = cell.paragraphs[0]
        p_value.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_value.paragraph_format.space_after = Pt(1)
        value_size = 18 if columns == 3 else 21
        set_run(p_value.add_run(metric["value"]), size=value_size, bold=True, color=BRAND_DARK)
        p_label = cell.add_paragraph()
        p_label.paragraph_format.space_after = Pt(0)
        set_run(
            p_label.add_run(metric["label"]),
            size=8.7 if columns == 3 else 9.5,
            bold=True,
            color=MUTED,
        )
    remove_table_cell_paragraph_spacing(table)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_after_hours_strip(doc, metrics):
    columns = len(metrics)
    if columns < 1:
        return
    widths = [CONTENT_WIDTH_DXA // columns] * columns
    widths[-1] += CONTENT_WIDTH_DXA - sum(widths)
    table = doc.add_table(rows=1, cols=columns)
    set_table_geometry(table, widths)
    set_table_borders(table, color=BRAND_GREEN, size=4)
    for index, metric in enumerate(metrics):
        cell = table.rows[0].cells[index]
        set_cell_shading(cell, BRAND_GREEN)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run(p.add_run(metric["value"]), size=13.5, bold=True, color=WHITE)
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run(p2.add_run(metric["label"]), size=7.8, bold=True, color=WHITE)
    remove_table_cell_paragraph_spacing(table)


def add_kpi_note(doc, note):
    if not note:
        return
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(5)
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.paragraph_format.left_indent = Inches(0.1)
    paragraph.paragraph_format.right_indent = Inches(0.05)
    set_paragraph_box(paragraph, fill="FFF7E5", border="D39A00")
    set_run(paragraph.add_run(note["title"] + ". "), size=9.3, bold=True, color="7A5A00")
    set_run(paragraph.add_run(note["text"]), size=9.3, color=INK)


def set_paragraph_box(paragraph, *, fill=PALE_GREEN, border=BRAND_GREEN):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)
    p_bdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), border)
    p_bdr.append(left)
    p_pr.append(p_bdr)


def add_highlight(doc, highlight):
    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(5)
    title.paragraph_format.space_after = Pt(2)
    set_run(title.add_run(highlight["title"]), size=11.5, bold=True, color=BRAND_DARK)

    summary = doc.add_paragraph()
    summary.paragraph_format.space_after = Pt(3)
    set_run(summary.add_run(highlight["summary"]), size=9.6, color=INK)

    quote = doc.add_paragraph()
    quote.paragraph_format.left_indent = Inches(0.11)
    quote.paragraph_format.right_indent = Inches(0.05)
    quote.paragraph_format.space_before = Pt(0)
    quote.paragraph_format.space_after = Pt(2)
    quote.paragraph_format.line_spacing = 1.05
    set_paragraph_box(quote)
    set_run(quote.add_run(highlight["quote"]), size=9.2, italic=True, color=INK)

    source = doc.add_paragraph()
    source.paragraph_format.space_after = Pt(1)
    set_run(source.add_run(highlight["source"]), size=8.2, bold=True, color=MUTED)


def add_feature(doc, feature):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.paragraph_format.keep_together = True
    set_run(paragraph.add_run(feature["title"] + ". "), size=10.5, bold=True, color=BRAND_DARK)
    set_run(paragraph.add_run(feature["description"]), size=10.5, color=INK)


def add_future_priorities(doc, priorities):
    label = doc.add_paragraph()
    label.paragraph_format.space_before = Pt(3)
    label.paragraph_format.space_after = Pt(3)
    set_run(label.add_run("NEAR-TERM DIRECTION"), size=8.2, bold=True, color=BRAND_GREEN)

    table = doc.add_table(rows=1, cols=2)
    set_table_geometry(table, [4680, 4680])
    set_table_borders(table, color="D7E4DC", size=5)
    for index, priority in enumerate(priorities):
        cell = table.rows[0].cells[index]
        set_cell_shading(cell, PALE_GREEN if index == 0 else PALE_GRAY)
        title = cell.paragraphs[0]
        title.paragraph_format.space_after = Pt(2)
        set_run(title.add_run(priority["title"]), size=11, bold=True, color=BRAND_DARK)
        description = cell.add_paragraph()
        description.paragraph_format.space_after = Pt(0)
        set_run(description.add_run(priority["description"]), size=9.4, color=INK)


def add_footer(section, report_date):
    footer = section.footer
    footer.is_linked_to_previous = False
    table = footer.add_table(rows=1, cols=2, width=Inches(6.5))
    set_table_geometry(table, [6500, 2860], indent_dxa=0)
    set_table_borders(table, color=WHITE, size=0)
    left = table.cell(0, 0).paragraphs[0]
    right = table.cell(0, 1).paragraphs[0]
    right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run(left.add_run("SpeakSport · Customer performance update"), size=8, color=MUTED)
    set_run(right.add_run(report_date), size=8, color=MUTED)


def configure_document(doc: Document):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.58)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.3)
    section.footer_distance = Inches(0.3)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = rgb(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.1

    for level, size, before, after in ((1, 16, 11, 5), (2, 12.5, 8, 3)):
        style = styles[f"Heading {level}"]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(BRAND_DARK)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def build(config_path: Path, output_path: Path):
    config = json.loads(config_path.read_text(encoding="utf-8"))
    operating_model = config.get("operating_model")
    if operating_model not in {"integrated", "non_integrated"}:
        raise ValueError(
            "update.json must set operating_model to 'integrated' or 'non_integrated'"
        )
    base_dir = config_path.parent
    logo_path = (base_dir / config["logo"]).resolve()

    doc = Document()
    configure_document(doc)
    section = doc.sections[0]
    add_footer(section, config["report_date"])

    logo_p = doc.add_paragraph()
    logo_p.paragraph_format.space_after = Pt(5)
    add_logo(logo_p, logo_path, 2.05)

    kicker = doc.add_paragraph()
    kicker.paragraph_format.space_after = Pt(0)
    set_run(kicker.add_run("CUSTOMER PERFORMANCE UPDATE"), size=8.5, bold=True, color=BRAND_GREEN)

    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(2)
    set_run(title.add_run(config["facility"]), size=25, bold=True, color=INK)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(7)
    set_run(
        subtitle.add_run(
            f"{config['dashboard_period']} · {config.get('service_model_label', '')}"
            f"{' · ' if config.get('service_model_label') else ''}Prepared {config['report_date']}"
        ),
        size=9.5,
        color=MUTED,
    )

    add_heading(doc, "Performance at a glance", 1)
    add_metric_grid(doc, config["headline_metrics"])
    add_kpi_note(doc, config.get("kpi_note"))
    label = doc.add_paragraph()
    label.paragraph_format.space_before = Pt(0)
    label.paragraph_format.space_after = Pt(3)
    set_run(label.add_run("AFTER-HOURS IMPACT"), size=8.2, bold=True, color=BRAND_GREEN)
    add_after_hours_strip(doc, config["after_hours_metrics"])

    add_heading(doc, "What has worked brilliantly", 1)
    intro = doc.add_paragraph()
    intro.paragraph_format.space_after = Pt(2)
    set_run(
        intro.add_run(
            f"A review of {config['call_review_count']} calls from {config['call_review_period']} "
            "showed three especially strong patterns:"
        ),
        size=9.6,
        color=MUTED,
    )
    for highlight in config["highlights"]:
        add_highlight(doc, highlight)

    doc.add_page_break()
    logo_p2 = doc.add_paragraph()
    logo_p2.paragraph_format.space_after = Pt(5)
    add_logo(logo_p2, logo_path, 1.65)

    add_heading(doc, "What we have rolled out", 1)
    intro2 = doc.add_paragraph()
    set_run(
        intro2.add_run(
            config.get(
                "features_intro",
                "We have continued strengthening the service around the moments that matter most to golfers and the pro shop.",
            )
        ),
        size=10.5,
        color=INK,
    )
    for feature in config["features"]:
        add_feature(doc, feature)

    add_heading(doc, "Building the next chapter together", 1)
    future = doc.add_paragraph()
    future.paragraph_format.space_after = Pt(7)
    set_run(future.add_run(config["future_note"]), size=10.5, color=INK)
    add_future_priorities(doc, config["future_priorities"])

    contact = doc.add_paragraph()
    contact.paragraph_format.space_before = Pt(2)
    contact.paragraph_format.space_after = Pt(10)
    contact.paragraph_format.left_indent = Inches(0.1)
    contact.paragraph_format.right_indent = Inches(0.1)
    set_paragraph_box(contact, fill=PALE_GREEN, border=BRAND_GREEN)
    set_run(contact.add_run("Have an idea for us? "), size=11, bold=True, color=BRAND_DARK)
    set_run(contact.add_run(config["contact_line"]), size=11, color=INK)

    note = doc.add_paragraph()
    note.paragraph_format.space_before = Pt(12)
    note.paragraph_format.space_after = Pt(0)
    set_run(note.add_run("About the figures and excerpts"), size=8.3, bold=True, color=MUTED)
    note2 = doc.add_paragraph()
    note2.paragraph_format.space_after = Pt(0)
    note2.paragraph_format.line_spacing = 1.0
    set_run(note2.add_run(config["method_note"]), size=7.8, color=MUTED)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.core_properties.title = f"SpeakSport performance update - {config['facility']}"
    doc.core_properties.subject = "Customer performance update"
    doc.core_properties.author = "SpeakSport"
    doc.core_properties.keywords = "SpeakSport, customer performance, golf operations"
    doc.save(output_path)


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("config", type=Path)
    parser.add_argument("output", type=Path)
    args = parser.parse_args()
    build(args.config, args.output)


if __name__ == "__main__":
    main()
